#!/usr/bin/env python3
"""
Virtual☆Paradise Local Diagnostic & Coding Agent
Cyberpunk-Themed Autonomous Offline Assistant with Full System Control & Dynamic Model Tiering
"""

import sys
import os
import subprocess
import json
import urllib.request
import urllib.error
import shutil
import readline
import time
import re
from typing import List, Dict, Any, Optional, Tuple

# Rich TUI Engine
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.markdown import Markdown
from rich.syntax import Syntax
from rich.text import Text
from rich import box

console = Console()

OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434")
DEFAULT_MODEL = os.environ.get("PARADISE_AGENT_MODEL", "qwen2.5-coder:3b")

# Execution Modes: "Auto-accept" (Default) or "Preview"
EXECUTION_MODE = "Auto-accept"
SHOW_THINKING = True

# Cyberpunk Palette Constants
CYAN = "#00f5d4"    # Neon Miku Cyan
GREEN = "#00ff88"   # Matrix Green
PINK = "#ffb7d5"    # Sakura Pink
PURPLE = "#b464ff"  # Neon Violet
YELLOW = "#ffe066"  # Cyberpunk Amber
RED = "#ff5287"     # Glitch Crimson
MUTED = "#708090"   # Terminal Slate

# Prompt Toolkit Interactive Engine (Dropdown Autocomplete & History)
try:
    from prompt_toolkit import PromptSession
    from prompt_toolkit.completion import Completer, Completion
    from prompt_toolkit.formatted_text import ANSI
    from prompt_toolkit.styles import Style
    from prompt_toolkit.history import FileHistory
    HAS_PROMPT_TOOLKIT = True
except ImportError:
    HAS_PROMPT_TOOLKIT = False

class SlashAndFileCompleter(Completer):
    """Provides dropdown autocomplete suggestions for '/' slash commands and '@' file tags."""
    COMMANDS = [
        ("/resume", "Session manager: resume, rename, or delete past conversations"),
        ("/plan", "Planning mode: generate phased diagnostic/coding plan before action"),
        ("/goal", "Goal mode: run autonomous multi-step execution until complete"),
        ("/model", "Switch local Ollama model (1.5b, 3b, 7b)"),
        ("/mode", "Toggle between Auto-accept and Preview approval modes"),
        ("/skills", "List all specialized expert skills discovered on machine"),
        ("/search", "Grep search across codebase using ripgrep: /search <query> [path]"),
        ("/find", "Find files by glob pattern using fd: /find <pattern> [dir]"),
        ("/health", "Hardware diagnostic health check (RAM, CPU, Cooler Boost, Disk)"),
        ("/thinking", "Toggle real-time diagnostic reasoning visibility (ON / OFF)"),
        ("/copy", "Copy last assistant response or output to Wayland clipboard"),
        ("/diff", "View git diff or repository status: /diff [path]"),
        ("/clear", "Reset active conversation memory context and start fresh"),
        ("/help", "Show detailed command reference table"),
        ("/exit", "Exit session (Sayonara)"),
        ("/quit", "Exit session (Sayonara)"),
    ]

    def get_completions(self, document, complete_event):
        text_before = document.text_before_cursor

        # 1. Slash command completion at start of line
        if text_before.startswith("/") and " " not in text_before:
            word = text_before.strip()
            for cmd, desc in self.COMMANDS:
                if cmd.startswith(word):
                    yield Completion(
                        cmd,
                        start_position=-len(text_before),
                        display=cmd,
                        display_meta=desc
                    )
            return

        # 2. @ File tag completion anywhere in prompt
        last_at = text_before.rfind("@")
        if last_at != -1:
            token = text_before[last_at:]
            if " " not in token:
                prefix = token[1:].strip()  # strip '@'
                search_dirs = [
                    os.getcwd(),
                    os.path.expanduser("~/Downloads"),
                    os.path.expanduser("~"),
                    os.path.expanduser("~/omarchy-virtual-paradise"),
                    os.path.expanduser("~/.config/hypr"),
                ]
                seen_files = set()
                count = 0
                for sdir in search_dirs:
                    if not os.path.exists(sdir) or count >= 30:
                        continue
                    try:
                        for entry in sorted(os.listdir(sdir)):
                            if entry.startswith(".git") or entry.startswith("__pycache__"):
                                continue
                            if not prefix or prefix.lower() in entry.lower():
                                if entry in seen_files:
                                    continue
                                seen_files.add(entry)
                                count += 1

                                full_p = os.path.join(sdir, entry)
                                is_dir = os.path.isdir(full_p)
                                if is_dir:
                                    meta = f"📁 Dir  | {sdir}"
                                else:
                                    try:
                                        sz = os.path.getsize(full_p)
                                        sz_str = f"{sz/1024:.1f} KB" if sz < 1024*1024 else f"{sz/(1024*1024):.1f} MB"
                                    except Exception:
                                        sz_str = ""
                                    meta = f"📄 {sz_str} | {sdir}"

                                yield Completion(
                                    f"@{entry}",
                                    start_position=-len(token),
                                    display=f"@{entry}",
                                    display_meta=meta
                                )
                                if count >= 30:
                                    break
                    except Exception:
                        pass

SKILL_DIRS = [
    os.path.expanduser("~/.agents/skills"),
    os.path.expanduser("~/Windows/skills"),
    os.path.expanduser("~/Windows/.agents/skills"),
    os.path.expanduser("~/Windows/agy-skills"),
    os.path.expanduser("~/.gemini/antigravity-cli/skills"),
    os.path.expanduser("~/.gemini/antigravity-cli/builtin/skills"),
    "/usr/share/omarchy/default/agents/skills"
]

SESSIONS_DIR = os.path.expanduser("~/.local/share/paradise-agent/sessions")

def get_session_filepath(session_id: str) -> str:
    return os.path.join(SESSIONS_DIR, f"{session_id}.json")

def save_session(session_id: str, messages: List[Dict[str, Any]], model: str, summary: str = ""):
    """Saves conversation state to JSON file."""
    try:
        os.makedirs(SESSIONS_DIR, exist_ok=True)
        fp = get_session_filepath(session_id)
        if not summary:
            for m in messages:
                if m.get("role") == "user" and m.get("content"):
                    raw_c = m["content"].split("[LANGUAGE MANDATE")[0].strip()
                    summary = raw_c.split("\n")[0][:65].strip()
                    break
        data = {
            "session_id": session_id,
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "model": model,
            "summary": summary or "Untitled Session",
            "messages": messages
        }
        with open(fp, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def list_saved_sessions() -> List[Dict[str, Any]]:
    """Returns sorted list of saved sessions (newest first)."""
    if not os.path.exists(SESSIONS_DIR):
        return []
    sessions = []
    for fn in os.listdir(SESSIONS_DIR):
        if fn.endswith(".json"):
            fp = os.path.join(SESSIONS_DIR, fn)
            try:
                with open(fp, "r", encoding="utf-8") as f:
                    d = json.load(f)
                    d["_mtime"] = os.path.getmtime(fp)
                    d["_filepath"] = fp
                    sessions.append(d)
            except Exception:
                pass
    sessions.sort(key=lambda x: x.get("_mtime", 0), reverse=True)
    return sessions

def load_session(session_id_or_number: str) -> Optional[Tuple[str, List[Dict[str, Any]], str]]:
    """Loads session by index number or session ID prefix. Returns (session_id, messages, model)."""
    sessions = list_saved_sessions()
    if not sessions:
        return None

    matched = None
    target = session_id_or_number.strip()
    if target.isdigit():
        idx = int(target) - 1
        if 0 <= idx < len(sessions):
            matched = sessions[idx]
    elif target.lower() in ("last", "latest", "default"):
        matched = sessions[0]

    if not matched:
        for s in sessions:
            sid = s.get("session_id", "")
            if sid == target or sid.startswith(target):
                matched = s
                break

    if matched:
        return matched.get("session_id", ""), matched.get("messages", []), matched.get("model", DEFAULT_MODEL)
    return None

def rename_session(session_id_or_number: str, new_title: str) -> bool:
    """Renames a saved session's title/summary."""
    sessions = list_saved_sessions()
    if not sessions:
        return False

    matched = None
    target = session_id_or_number.strip()
    if target.isdigit():
        idx = int(target) - 1
        if 0 <= idx < len(sessions):
            matched = sessions[idx]
    else:
        for s in sessions:
            sid = s.get("session_id", "")
            if sid == target or sid.startswith(target):
                matched = s
                break

    if not matched:
        return False

    fp = matched.get("_filepath")
    if not fp or not os.path.exists(fp):
        return False

    try:
        data_to_write = {
            "session_id": matched.get("session_id"),
            "updated_at": matched.get("updated_at"),
            "model": matched.get("model"),
            "summary": new_title.strip(),
            "messages": matched.get("messages", [])
        }
        with open(fp, "w", encoding="utf-8") as f:
            json.dump(data_to_write, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False

def delete_session(session_id_or_number: str) -> bool:
    """Deletes a saved session from disk."""
    sessions = list_saved_sessions()
    if not sessions:
        return False

    matched = None
    target = session_id_or_number.strip()
    if target.isdigit():
        idx = int(target) - 1
        if 0 <= idx < len(sessions):
            matched = sessions[idx]
    else:
        for s in sessions:
            sid = s.get("session_id", "")
            if sid == target or sid.startswith(target):
                matched = s
                break

    if not matched:
        return False

    fp = matched.get("_filepath")
    if fp and os.path.exists(fp):
        try:
            os.remove(fp)
            return True
        except Exception:
            return False
    return False

def render_restored_session(messages: List[Dict[str, Any]]):
    """Renders the restored chat transcript into the terminal."""
    if not messages:
        return

    user_turns = sum(1 for m in messages if m.get("role") == "user")
    if user_turns == 0:
        return

    console.print(Panel(
        f"[bold {CYAN}]Replaying Previous Chat Transcript ({len(messages)} messages)[/]",
        border_style=CYAN,
        box=box.ROUNDED,
        padding=(0, 1)
    ))

    for m in messages:
        role = m.get("role")
        if role == "system":
            continue

        elif role == "user":
            raw_c = m.get("content", "")
            clean_c = raw_c.split("[LANGUAGE MANDATE")[0].strip()
            console.print(f"\n\033[38;2;0;245;212m\033[1myou ❯\033[0m {clean_c}")

        elif role == "assistant":
            tool_calls = m.get("tool_calls", [])
            for tc in tool_calls:
                fn = tc.get("function", {})
                fn_name = fn.get("name", "action")
                fn_args = fn.get("arguments", {})
                if isinstance(fn_args, str):
                    try:
                        fn_args = json.loads(fn_args)
                    except Exception:
                        pass
                arg_repr = json.dumps(fn_args, ensure_ascii=False) if isinstance(fn_args, dict) else str(fn_args)
                console.print(Panel(
                    f"{fn_name}({arg_repr})",
                    title=f"[bold {YELLOW}]⚙ Tool Execution[/]",
                    border_style=MUTED,
                    box=box.ROUNDED,
                    padding=(0, 1)
                ))

            content = m.get("content", "")
            if content:
                console.print(f"\n[bold {PINK}]paradise-agent ❯[/]")
                console.print(Markdown(content))

        elif role == "tool":
            raw_out = m.get("content", "")
            clean_out = raw_out.split("[SYSTEM MANDATE")[0].strip()
            lines = clean_out.split("\n")
            preview = "\n".join(lines[:5])
            if len(lines) > 5:
                preview += f"\n... ({len(lines) - 5} more lines)"
            console.print(Panel(
                preview,
                title=f"[bold {GREEN}]📄 Tool Result ({m.get('name', 'tool')})[/]",
                border_style=MUTED,
                box=box.ROUNDED,
                padding=(0, 1)
            ))

    console.print(Panel(
        f"[bold {GREEN}]✔ Conversation context restored. Continue typing below:[/]",
        border_style=GREEN,
        box=box.ROUNDED,
        padding=(0, 1)
    ))

def handle_resume_command(arg: str = "") -> Optional[Tuple[str, List[Dict[str, Any]], str]]:
    """Interactive /resume handler with resume, rename, and delete operations."""
    sessions = list_saved_sessions()
    if not sessions:
        console.print(f"[bold {YELLOW}]No saved sessions found in {SESSIONS_DIR}[/]")
        return None

    if arg:
        parts = arg.split(maxsplit=2)
        subcmd = parts[0].lower()
        if subcmd in ("rename", "r", "mv", "name"):
            if len(parts) >= 3:
                target_id = parts[1]
                new_title = parts[2]
                if rename_session(target_id, new_title):
                    console.print(f"[bold {GREEN}]✔ Successfully renamed session '{target_id}' to: '{new_title}'[/]")
                else:
                    console.print(f"[bold {RED}]Could not find session '{target_id}' to rename.[/]")
            else:
                console.print(f"[bold {YELLOW}]Usage: /resume rename <#|session_id> <new_title>[/]")
            return None

        elif subcmd in ("delete", "del", "rm", "d"):
            if len(parts) >= 2:
                target_id = parts[1]
                if delete_session(target_id):
                    console.print(f"[bold {GREEN}]✔ Successfully deleted session '{target_id}'[/]")
                else:
                    console.print(f"[bold {RED}]Could not find session '{target_id}' to delete.[/]")
            else:
                console.print(f"[bold {YELLOW}]Usage: /resume delete <#|session_id>[/]")
            return None

        res = load_session(arg)
        if res:
            return res
        console.print(f"[bold {RED}]Session '{arg}' not found.[/]")
        return None

    while True:
        sessions = list_saved_sessions()
        if not sessions:
            console.print(f"[bold {YELLOW}]No saved sessions remaining.[/]")
            return None

        table = Table(title=f"Saved Conversations ({len(sessions)} available)", border_style=CYAN, box=box.ROUNDED)
        table.add_column("#", style=f"bold {YELLOW}", width=4)
        table.add_column("Session ID", style=f"bold {CYAN}", width=25)
        table.add_column("Last Updated", style=f"{MUTED}", width=19)
        table.add_column("Model", style=f"{GREEN}", width=16)
        table.add_column("Title / Summary", style=f"{PINK}")

        for idx, s in enumerate(sessions[:15]):
            sid = s.get("session_id", "")
            upd = s.get("updated_at", "")
            mod = s.get("model", "")
            summ = s.get("summary", "")
            table.add_row(str(idx + 1), sid, upd, mod, summ)

        console.print(table)
        console.print(
            f"[bold {CYAN}]Options:[/] Enter [1-{min(len(sessions), 15)}] to resume  |  "
            f"[bold {YELLOW}]r <#> <name>[/] to rename  |  "
            f"[bold {RED}]d <#>[/] to delete  |  "
            f"[bold {MUTED}]c[/] to cancel"
        )

        try:
            choice = input(f"\033[38;2;255;224;102mAction [Enter=1 / r # <name> / d # / c]: \033[0m").strip()
            if not choice:
                choice = "1"
            if choice.lower() in ("c", "cancel", "q", "quit"):
                console.print(f"[dim]Resume cancelled.[/]")
                return None

            if choice.lower().startswith("r ") or choice.lower().startswith("rename "):
                r_parts = choice.split(maxsplit=2)
                if len(r_parts) >= 3:
                    if rename_session(r_parts[1], r_parts[2]):
                        console.print(f"[bold {GREEN}]✔ Renamed session #{r_parts[1]} to: '{r_parts[2]}'[/]")
                    else:
                        console.print(f"[bold {RED}]Failed to rename session #{r_parts[1]}[/]")
                else:
                    console.print(f"[bold {YELLOW}]Usage: r <#> <new_name>[/]")
                continue

            if choice.lower().startswith("d ") or choice.lower().startswith("del ") or choice.lower().startswith("rm ") or choice.lower().startswith("delete "):
                d_parts = choice.split(maxsplit=1)
                if len(d_parts) >= 2:
                    if delete_session(d_parts[1]):
                        console.print(f"[bold {GREEN}]✔ Deleted session #{d_parts[1]}[/]")
                    else:
                        console.print(f"[bold {RED}]Failed to delete session #{d_parts[1]}[/]")
                else:
                    console.print(f"[bold {YELLOW}]Usage: d <#>[/]")
                continue

            loaded = load_session(choice)
            if loaded:
                return loaded
            console.print(f"[bold {RED}]Invalid session selection: '{choice}'[/]")
        except (KeyboardInterrupt, EOFError):
            console.print("")
            return None

def get_system_memory_info() -> Dict[str, float]:
    """Returns system memory metrics in GiB from /proc/meminfo."""
    info = {"total": 16.0, "available": 8.0, "free": 2.0}
    try:
        with open("/proc/meminfo", "r") as f:
            for line in f:
                if ":" in line:
                    k, v = line.split(":", 1)
                    k = k.strip()
                    val = v.strip().split()[0]
                    if k == "MemTotal":
                        info["total"] = round(int(val) / (1024 * 1024), 2)
                    elif k == "MemAvailable":
                        info["available"] = round(int(val) / (1024 * 1024), 2)
                    elif k == "MemFree":
                        info["free"] = round(int(val) / (1024 * 1024), 2)
    except Exception:
        pass
    return info

def get_installed_models() -> List[str]:
    """Queries installed local models from Ollama API."""
    try:
        req = urllib.request.Request(f"{OLLAMA_HOST}/api/tags")
        with urllib.request.urlopen(req, timeout=3) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return [m.get("name", "") for m in data.get("models", [])]
    except Exception:
        return [DEFAULT_MODEL]

def detect_optimal_model(requested_model: Optional[str] = None) -> Tuple[str, str]:
    """
    Intelligently inspects available system RAM and chooses the best installed model:
    - Available RAM >= 6.5 GiB: qwen2.5-coder:7b (High Capability Tier)
    - Available RAM 3.0 - 6.5 GiB: qwen2.5-coder:3b (Sweet Spot Tier)
    - Available RAM < 3.0 GiB: qwen2.5-coder:1.5b (Ultra-lightweight Tier)
    """
    if requested_model:
        return requested_model, f"manual override: {requested_model}"

    env_override = os.environ.get("PARADISE_AGENT_MODEL")
    if env_override:
        return env_override, f"environment variable PARADISE_AGENT_MODEL ({env_override})"

    mem = get_system_memory_info()
    avail = mem["available"]
    installed = get_installed_models()

    has_7b = any("7b" in m for m in installed)
    has_3b = any("3b" in m for m in installed)
    has_1_5b = any("1.5b" in m for m in installed)

    # 3B is the optimal high-speed sweet spot for local CPU execution (~2-3s response)
    if has_3b and avail >= 3.0:
        chosen = [m for m in installed if "3b" in m][0]
        return chosen, f"Available RAM: {avail} GiB (Tier: 3B Fast & Optimal)"
    elif avail < 3.0 and has_1_5b:
        chosen = [m for m in installed if "1.5b" in m][0]
        return chosen, f"Available RAM: {avail} GiB (Tier: 1.5B Ultra-light)"
    elif has_7b:
        chosen = [m for m in installed if "7b" in m][0]
        return chosen, f"Available RAM: {avail} GiB (Tier: 7B)"
    elif has_3b:
        chosen = [m for m in installed if "3b" in m][0]
        return chosen, f"Available RAM: {avail} GiB (Tier: 3B)"
    elif installed:
        return installed[0], f"Installed model: {installed[0]}"
    return DEFAULT_MODEL, f"Default offline ({DEFAULT_MODEL})"

def discover_skills() -> Dict[str, Dict[str, str]]:
    """Discovers all specialized skills across known repository directories."""
    skills = {}
    for base in SKILL_DIRS:
        if not os.path.exists(base):
            continue
        try:
            items = sorted(os.listdir(base))
        except Exception:
            continue
        for item in items:
            skill_path = os.path.join(base, item, "SKILL.md")
            if os.path.isfile(skill_path) and item not in skills:
                desc = ""
                try:
                    with open(skill_path, "r", encoding="utf-8", errors="replace") as f:
                        content = f.read()
                    if content.startswith("---"):
                        parts = content.split("---", 2)
                        if len(parts) >= 3:
                            for line in parts[1].split("\n"):
                                if line.strip().startswith("description:"):
                                    desc = line.split("description:", 1)[1].strip(" >-\t\r\n")
                                elif desc and (line.startswith("  ") or line.startswith("\t")):
                                    desc += " " + line.strip()
                    if not desc:
                        desc = f"Specialized expert instructions for {item}"
                except Exception:
                    desc = f"Specialized expert instructions for {item}"
                skills[item] = {
                    "name": item,
                    "path": skill_path,
                    "dir": os.path.join(base, item),
                    "description": desc.strip()
                }
    return skills

def build_system_prompt() -> str:
    """Builds the comprehensive system prompt for the autonomous agent."""
    theme = "Virtual Paradise"
    try:
        p = os.path.expanduser("~/.config/omarchy/current/theme.name")
        if os.path.exists(p):
            with open(p) as f:
                t = f.read().strip()
                if t:
                    theme = t
    except Exception:
        pass

    user = os.environ.get("USER", "doe")
    shell = os.environ.get("SHELL", "/usr/bin/zsh")

    skills_map = discover_skills()
    skills_lines = []
    for name, s in sorted(skills_map.items()):
        skills_lines.append(f"  • {name}: {s['description']}")
    skills_block = "\n".join(skills_lines)

    return f"""You are Paradise Agent — an autonomous local AI pair-programming, system diagnostic, and engineering assistant on Arch Linux (Hyprland/Wayland, Omarchy theme: "{theme}").
User: {user} | Shell: {shell} | Terminal: Ghostty | Desktop: Hyprland
Key paths: ~/omarchy-virtual-paradise (theme/dotfiles), ~/Windows (shared data/.bat scripts), ~/Downloads

=== AVAILABLE EXPERT SKILLS (Progressive Disclosure) ===
{skills_block}

=== CAPABILITIES & ROOT PERMISSIONS ===
- You have UNRESTRICTED FULL ROOT / SUDO permissions across the entire system.
- You can READ, WRITE, EDIT, DELETE, COPY, and MOVE ANY file anywhere on the computer without restrictions.
- You can execute ANY bash/system command (including systemctl, pacman, isw, ip, hyprctl, journalctl).
- NEVER refuse a command with "I cannot access", "I do not have permission", "I cannot help".

=== AGY AUTONOMOUS PAIR-PROGRAMMING DIRECTIVES ===
1. BIAS TOWARDS ACTION: Execute tools directly to inspect, code, test, and diagnose. Do not emit hypothetical markdown bash blocks asking the user to run commands manually.
2. PROACTIVE GROUNDING (VERIFY BEFORE CONCLUDING):
   - Never assume a file exists or doesn't exist. Use `find_by_name`, `list_dir`, or `read_file` to verify filesystem reality first.
   - Use `grep_search` to find symbols, functions, configurations, and errors before editing.
3. SURGICAL CODE EDITS & INTEGRITY:
   - Prefer `edit_file` to replace exact contiguous blocks rather than overwriting whole files.
   - Preserve existing comments, docstrings, formatting, and structures unless explicitly asked to modify them.
4. PROGRESSIVE DISCLOSURE FOR SKILLS:
   - When a task involves specialized areas (e.g. Hyprland/desktop config -> omarchy, crash analysis -> diagnose-crash, Vietnamese documents -> vn-officecli), call `load_skill` to retrieve the expert runbook.
5. STRICT BILINGUAL CONSISTENCY:
   - DEFAULT LANGUAGE: ENGLISH. By default, respond, explain, and write diagnostic reports in clear, professional ENGLISH.
   - EXCEPTION: If and only if the user inputs their query in VIETNAMESE, formulate your entire response in natural, fluent, polite VIETNAMESE.
   - STRICT PROHIBITION: NEVER reply in French, German, Spanish, Chinese, or any other language. If user writes Vietnamese, reply ONLY in Vietnamese.

=== TOOL CALL EXAMPLES ===
User: format file tờ trình trong mục download
Assistant: {{"name": "execute_bash", "arguments": {{"command": "/home/{user}/.local/lib/paradise-venv/bin/python3 /home/{user}/.local/bin/format-docx-vn.py /home/{user}/Downloads"}}}}

User: xóa skill windows đi
Assistant: {{"name": "delete_skill", "arguments": {{"skill_name": "windows"}}}}

User: tìm hàm lerp trong repo theme
Assistant: {{"name": "grep_search", "arguments": {{"query": "def lerp", "search_path": "/home/{user}/omarchy-virtual-paradise"}}}}

User: tìm file to_trinh
Assistant: {{"name": "find_by_name", "arguments": {{"pattern": "*to_trinh*", "search_directory": "/home/{user}"}}}}

User: trong thư mục Downloads có gì
Assistant: {{"name": "list_dir", "arguments": {{"directory_path": "/home/{user}/Downloads"}}}}

User: giúp mình xóa file formatted
Assistant: {{"name": "delete_file", "arguments": {{"path": "/home/{user}/Downloads/to_trinh_mau_formatted.docx"}}}}

User: nội dung file này có gì
Assistant: {{"name": "read_file", "arguments": {{"path": "/home/{user}/Downloads/to_trinh_mau.docx"}}}}

User: bật cooler boost
Assistant: {{"name": "execute_bash", "arguments": {{"command": "/home/{user}/.local/bin/toggle_cooler_boost.sh on"}}}}

User: check system health
Assistant: {{"name": "get_system_health", "arguments": {{}}}}
"""

TOOLS_SPEC = [
    {
        "type": "function",
        "function": {
            "name": "execute_bash",
            "description": "Execute a bash command on the local system with full root/sudo permissions.",
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "bash command string"
                    }
                },
                "required": ["command"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read any file or directory across the entire filesystem (supports Word .docx, configs, logs, code, text).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "path to file or directory"
                    },
                    "max_lines": {
                        "type": "integer",
                        "description": "maximum lines to return (default: 200)"
                    }
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Write or overwrite text content to any file path on the system with full root/sudo permissions.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "target file path"
                    },
                    "content": {
                        "type": "string",
                        "description": "full text content to write"
                    }
                },
                "required": ["path", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "edit_file",
            "description": "Replace specific text in an existing file (works on configs, code, scripts, and Word .docx documents).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "file path"
                    },
                    "target_text": {
                        "type": "string",
                        "description": "exact text snippet to replace"
                    },
                    "replacement_text": {
                        "type": "string",
                        "description": "new replacement text"
                    }
                },
                "required": ["path", "target_text", "replacement_text"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_file",
            "description": "Delete any file or directory anywhere on the system (uses root elevation if required).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "file or directory path to delete"
                    }
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "copy_file",
            "description": "Copy a file or directory to a destination path anywhere on the system.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "source path"
                    },
                    "destination": {
                        "type": "string",
                        "description": "destination path"
                    }
                },
                "required": ["source", "destination"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "move_file",
            "description": "Move or rename a file or directory anywhere on the system.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {
                        "type": "string",
                        "description": "source path"
                    },
                    "destination": {
                        "type": "string",
                        "description": "destination path"
                    }
                },
                "required": ["source", "destination"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_system_health",
            "description": "Inspect CPU, RAM, Disk, Battery, top processes, and failed system services.",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "load_skill",
            "description": "Load expert instructions for a specialized skill (e.g. omarchy, diagnose-crash, vn-officecli).",
            "parameters": {
                "type": "object",
                "properties": {
                    "skill_name": {
                        "type": "string",
                        "description": "skill name"
                    }
                },
                "required": ["skill_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_skill",
            "description": "Delete and remove an installed skill from the system by its name (e.g. windows, omarchy, vn-officecli).",
            "parameters": {
                "type": "object",
                "properties": {
                    "skill_name": {
                        "type": "string",
                        "description": "the exact name of the skill to delete"
                    }
                },
                "required": ["skill_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "grep_search",
            "description": "Fast ripgrep search for text patterns, functions, configs, or errors across codebase.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "text pattern to search"
                    },
                    "search_path": {
                        "type": "string",
                        "description": "directory or file path to search in (default: .)"
                    },
                    "case_insensitive": {
                        "type": "boolean",
                        "description": "whether search is case insensitive (default: true)"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "find_by_name",
            "description": "Fast fd search for files and folders matching pattern across directory tree.",
            "parameters": {
                "type": "object",
                "properties": {
                    "pattern": {
                        "type": "string",
                        "description": "glob pattern (e.g. *.py, to_trinh*, hyprland.conf)"
                    },
                    "search_directory": {
                        "type": "string",
                        "description": "root directory to search (default: .)"
                    }
                },
                "required": ["pattern"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_dir",
            "description": "List files and directories with formatted sizes and icons.",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory_path": {
                        "type": "string",
                        "description": "path to list (default: .)"
                    }
                }
            }
        }
    }
]

def unwrap_tool_arg(val: Any) -> Any:
    """Recursively unwraps tool arguments from LLM schema wrappers."""
    if isinstance(val, dict):
        if "value" in val and len(val) == 1:
            return unwrap_tool_arg(val["value"])
        if "description" in val and "type" in val:
            return ""
        if len(val) == 1 and any(k in val for k in ("path", "command", "skill_name", "target_text", "replacement_text", "source", "destination")):
            for k in ("path", "command", "skill_name", "target_text", "replacement_text", "source", "destination"):
                if k in val:
                    return unwrap_tool_arg(val[k])
    return val

ALL_VALID_TOOLS = {
    "execute_bash",
    "read_file",
    "write_file",
    "edit_file",
    "delete_file",
    "copy_file",
    "move_file",
    "grep_search",
    "find_by_name",
    "list_dir",
    "load_skill",
    "delete_skill",
    "get_system_health",
}

TOOL_ALIASES = {
    "bash": "execute_bash",
    "run_command": "execute_bash",
    "sh": "execute_bash",
    "terminal": "execute_bash",
    "exec": "execute_bash",
    "shell": "execute_bash",
    "command": "execute_bash",
    "readfile": "read_file",
    "view_file": "read_file",
    "cat": "read_file",
    "open_file": "read_file",
    "writefile": "write_file",
    "write_to_file": "write_file",
    "create_file": "write_file",
    "editfile": "edit_file",
    "replace_file_content": "edit_file",
    "update_file": "edit_file",
    "grep": "grep_search",
    "ripgrep": "grep_search",
    "search": "grep_search",
    "find_text": "grep_search",
    "find": "find_by_name",
    "locate": "find_by_name",
    "search_files": "find_by_name",
    "ls": "list_dir",
    "listdir": "list_dir",
    "dir": "list_dir",
    "list_directory": "list_dir",
    "delete": "delete_file",
    "rm": "delete_file",
    "remove_file": "delete_file",
    "copy": "copy_file",
    "cp": "copy_file",
    "move": "move_file",
    "mv": "move_file",
    "rename": "move_file",
    "deleteskill": "delete_skill",
    "remove_skill": "delete_skill",
    "loadskill": "load_skill",
    "health": "get_system_health",
    "system_health": "get_system_health",
    "check_health": "get_system_health",
}

def normalize_tool_call(fn_name: str, fn_args: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    """Normalizes tool name and argument keys to prevent LLM schema hallucinations."""
    clean_name = fn_name.lower().replace("-", "_").strip()
    canonical_name = TOOL_ALIASES.get(clean_name, clean_name)

    normalized_args = {}
    for k, v in fn_args.items():
        val = unwrap_tool_arg(v)
        k_norm = k.lower().replace("-", "_").strip()
        if k_norm in ("file_path", "filename", "filepath", "target_file", "file"):
            normalized_args["path"] = val
        elif k_norm in ("cmd", "command_line", "shell_cmd"):
            normalized_args["command"] = val
        elif k_norm in ("search_term", "pattern", "text") and canonical_name == "grep_search":
            normalized_args["query"] = val
        elif k_norm in ("dir", "directory", "folder", "folder_path") and canonical_name == "list_dir":
            normalized_args["directory_path"] = val
        elif k_norm in ("dir", "directory", "folder", "search_dir") and canonical_name in ("find_by_name", "grep_search"):
            if canonical_name == "find_by_name":
                normalized_args["search_directory"] = val
            else:
                normalized_args["search_path"] = val
        elif k_norm in ("name", "skill") and canonical_name in ("load_skill", "delete_skill"):
            normalized_args["skill_name"] = val
        elif k_norm in ("src", "from"):
            normalized_args["source"] = val
        elif k_norm in ("dest", "dst", "to"):
            normalized_args["destination"] = val
        else:
            normalized_args[k] = val

    return canonical_name, normalized_args

def prompt_preview_action(tool_name: str, action_summary: str, details: str = "") -> bool:
    """Prompt user in Preview mode before executing an action."""
    global EXECUTION_MODE
    if EXECUTION_MODE == "Auto-accept":
        return True

    preview_body = (
        f"[bold {YELLOW}]Tool:[/] [bold {CYAN}]{tool_name}[/]\n"
        f"[bold {YELLOW}]Action:[/] {action_summary}\n"
    )
    if details:
        preview_body += f"[bold {MUTED}]Details:[/] {details}\n"

    console.print(Panel(
        preview_body.strip(),
        title=f"[bold {YELLOW}]⚠️ Action Preview Approval Required[/]",
        border_style=YELLOW,
        box=box.ROUNDED
    ))

    try:
        choice = input(f"\033[38;2;255;224;102mExecute this action? [Y/n/auto]: \033[0m").strip().lower()
        if choice in ("a", "auto", "auto-accept"):
            EXECUTION_MODE = "Auto-accept"
            console.print(f"[bold {GREEN}]Switched execution mode to Auto-accept.[/]")
            return True
        elif choice in ("n", "no", "cancel"):
            console.print(f"[bold {RED}]Action cancelled by user.[/]")
            return False
        return True
    except (KeyboardInterrupt, EOFError):
        console.print(f"\n[bold {RED}]Action cancelled.[/]")
        return False

def tool_execute_bash(command: Any) -> str:
    """Executes a bash command with full root/sudo privileges, auto-elevation, and non-interactive safeguards."""
    command = str(unwrap_tool_arg(command)).strip()
    if not command:
        return "[Error: Empty command]"

    if command.startswith("delete_skill"):
        parts = command.split()
        sk = parts[1] if len(parts) > 1 else "windows"
        return tool_delete_skill(sk)
    elif command.startswith("load_skill"):
        parts = command.split()
        sk = parts[1] if len(parts) > 1 else "windows"
        return tool_load_skill(sk)

    # Automatically add --noconfirm for pacman package installations/removals to prevent hanging
    if "pacman" in command and "--noconfirm" not in command:
        command = re.sub(r'\bpacman\s+(-[SsyuURm]+)\b', r'pacman \1 --noconfirm', command)

    if not prompt_preview_action("execute_bash", command):
        return "[Action Cancelled: User declined execution in Preview mode]"

    cmd_env = os.environ.copy()
    cmd_env["DEBIAN_FRONTEND"] = "noninteractive"
    cmd_env["PAGER"] = "cat"
    cmd_env["CI"] = "1"

    try:
        p = subprocess.run(
            ["bash", "-c", command],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=60,
            env=cmd_env
        )
        out = p.stdout.strip()
        err = p.stderr.strip()

        # Auto-elevate with sudo if permission denied and not already sudo
        if p.returncode != 0 and ("permission denied" in err.lower() or "operation not permitted" in err.lower()) and not command.startswith("sudo"):
            sudo_cmd = f"sudo bash -c '{command}'"
            p_sudo = subprocess.run(
                ["bash", "-c", sudo_cmd],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=60,
                env=cmd_env
            )
            out = p_sudo.stdout.strip()
            err = p_sudo.stderr.strip()
            p = p_sudo

        combined = ""
        if p.returncode != 0:
            combined += f"[Exit code: {p.returncode}]\n"
        if out:
            combined += out
        if err:
            combined += ("\n[STDERR]:\n" if combined else "") + err
        if not combined:
            combined = "[Command executed successfully with no output]"

        # Intelligent head/tail truncation if output is massive
        if len(combined) > 6000:
            omitted = len(combined) - 5500
            combined = combined[:3000] + f"\n... [Omitted {omitted} characters of verbose output] ...\n" + combined[-2500:]

        return combined
    except subprocess.TimeoutExpired:
        return "[Error: Command execution timed out after 60 seconds]"
    except Exception as e:
        return f"[Execution Error: {e}]"

def tool_read_file(path: Any, start_line: Optional[int] = None, end_line: Optional[int] = None, max_lines: int = 200) -> str:
    """Reads any file or directory with native Word .docx and PDF support, line slicing, and root fallback."""
    path = str(unwrap_tool_arg(path)).strip()
    if not path or path in ("{}", "None", "''", '""'):
        path = os.path.expanduser("~/Downloads")
    expanded_path = os.path.expanduser(path)

    # Auto-resolve relative or basename filenames to standard dirs
    if not os.path.exists(expanded_path):
        candidates = [
            os.path.join(os.path.expanduser("~/Downloads"), path),
            os.path.join(os.path.expanduser("~"), path),
            os.path.join(os.path.expanduser("~/omarchy-virtual-paradise"), path),
            os.path.join(os.path.expanduser("~/Windows"), path),
            os.path.join(os.path.expanduser("~/.config/hypr"), path),
        ]
        for c in candidates:
            if os.path.exists(c):
                expanded_path = c
                path = c
                break
        if not os.path.exists(expanded_path):
            return f"[Error: File or directory '{path}' not found]"

    if os.path.isdir(expanded_path):
        return tool_list_dir(expanded_path)

    # Native extraction for PDF documents
    if expanded_path.lower().endswith(".pdf"):
        try:
            cmd = ["pdftotext", "-layout", "-f", "1", "-l", "25", expanded_path, "-"]
            out = subprocess.check_output(cmd, text=True, timeout=8)
            if not out.strip():
                out = subprocess.check_output(["pdftotext", "-f", "1", "-l", "25", expanded_path, "-"] , text=True, timeout=8)
            return f"[PDF Document Content '{os.path.basename(path)}']:\n{out.strip()[:6000]}"
        except Exception as e:
            return f"[Error extracting PDF document: {e}]"

    # Native extraction for Word .docx
    if expanded_path.lower().endswith(".docx"):
        try:
            venv_py = os.path.expanduser("~/.local/lib/paradise-venv/bin/python3")
            py_bin = venv_py if os.path.exists(venv_py) else "python3"
            cmd = [
                py_bin, "-c",
                "import sys; from docx import Document; doc=Document(sys.argv[1]); "
                "paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]; "
                "print('\\n'.join(paras[:120]))",
                expanded_path
            ]
            out = subprocess.check_output(cmd, text=True, timeout=8)
            return f"[Word Document Content '{os.path.basename(path)}']:\n{out.strip()}"
        except Exception as e:
            return f"[Error extracting Word document: {e}]"

    try:
        with open(expanded_path, "r", encoding="utf-8", errors="replace") as f:
            all_lines = f.readlines()
        
        if start_line is not None and end_line is not None:
            s_idx = max(0, int(start_line) - 1)
            e_idx = min(len(all_lines), int(end_line))
            sliced = all_lines[s_idx:e_idx]
            numbered = [f"{s_idx + 1 + i:4d}: {l}" for i, l in enumerate(sliced)]
            return "".join(numbered)
        else:
            return "".join(all_lines[:max_lines])
    except PermissionError:
        try:
            out = subprocess.check_output(["sudo", "head", f"-n{max_lines}", expanded_path], text=True, timeout=5)
            return out
        except Exception as e:
            return f"[Permission Error reading '{path}': {e}]"
    except Exception as e:
        return f"[Error reading file: {e}]"

def tool_list_dir(directory_path: Any = None) -> str:
    """Lists files and folders in directory with human-readable sizes and icons."""
    dp = str(unwrap_tool_arg(directory_path or ".")).strip()
    if not dp or dp in ("None", "''"):
        dp = "."
    expanded_dp = os.path.expanduser(dp)
    if not os.path.exists(expanded_dp):
        return f"[Error: Directory '{dp}' not found]"

    try:
        entries = sorted(os.listdir(expanded_dp))
        out = []
        for e in entries[:60]:
            fp = os.path.join(expanded_dp, e)
            if os.path.isdir(fp):
                out.append(f"📁 {e}/")
            else:
                sz = os.path.getsize(fp)
                if sz < 1024:
                    sz_str = f"{sz} B"
                elif sz < 1024 * 1024:
                    sz_str = f"{sz // 1024} KB"
                else:
                    sz_str = f"{sz // (1024*1024)} MB"
                out.append(f"📄 {e} ({sz_str})")
        if len(entries) > 60:
            out.append(f"... ({len(entries) - 60} more items)")
        return f"[Directory listing of '{dp}']:\n" + "\n".join(out)
    except PermissionError:
        try:
            return subprocess.check_output(["sudo", "ls", "-lah", expanded_dp], text=True, timeout=5)[:3000]
        except Exception as e:
            return f"[Permission error listing '{dp}': {e}]"
    except Exception as e:
        return f"[Error listing directory: {e}]"

def tool_grep_search(query: Any, search_path: Any = None, is_regex: bool = False, case_insensitive: bool = True) -> str:
    """Fast ripgrep text search across directory or files."""
    query = str(unwrap_tool_arg(query)).strip()
    sp = str(unwrap_tool_arg(search_path or ".")).strip()
    if not sp or sp in ("None", "''"):
        sp = "."
    expanded_path = os.path.expanduser(sp)
    if not os.path.exists(expanded_path):
        return f"[Error: Search path '{sp}' does not exist]"

    cmd = ["rg", "-n", "--max-count=30", "--max-columns=120"]
    if case_insensitive:
        cmd.append("-i")
    if not is_regex:
        cmd.append("-F")
    cmd.extend([query, expanded_path])

    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
        lines = p.stdout.strip().split("\n")
        if not p.stdout.strip():
            return f"[No matches found for '{query}' in '{sp}']"
        out = "\n".join(lines[:35])
        if len(lines) > 35:
            out += f"\n... ({len(lines) - 35} more matches truncated)"
        return out
    except Exception as e:
        return f"[Grep search error: {e}]"

def tool_find_by_name(pattern: Any, search_directory: Any = None) -> str:
    """Fast fd search for files and folders by name pattern."""
    pat = str(unwrap_tool_arg(pattern)).strip()
    sd = str(unwrap_tool_arg(search_directory or ".")).strip()
    if not sd or sd in ("None", "''"):
        sd = "."
    expanded_sd = os.path.expanduser(sd)
    if not os.path.exists(expanded_sd):
        return f"[Error: Directory '{sd}' does not exist]"

    cmd = ["fd", "--hidden", "--exclude", ".git", "--max-results=40", pat, expanded_sd]
    try:
        p = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
        out = p.stdout.strip()
        if not out:
            return f"[No files found matching '{pat}' in '{sd}']"
        return out
    except Exception as e:
        return f"[Find error: {e}]"

def tool_write_file(path: Any, content: Any) -> str:
    """Writes or overwrites text to any file path, elevating to sudo if required."""
    path = str(unwrap_tool_arg(path)).strip()
    content = str(unwrap_tool_arg(content))
    if not path:
        return "[Error: Missing target file path]"
    expanded_path = os.path.expanduser(path)

    if not prompt_preview_action("write_file", f"Write to '{expanded_path}' ({len(content)} bytes)"):
        return "[Action Cancelled: User declined write operation]"

    try:
        os.makedirs(os.path.dirname(os.path.abspath(expanded_path)), exist_ok=True)
        with open(expanded_path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"[Success: File '{path}' written successfully ({len(content)} bytes)]"
    except PermissionError:
        try:
            p = subprocess.run(["sudo", "tee", expanded_path], input=content, text=True, capture_output=True, timeout=10)
            if p.returncode == 0:
                return f"[Success: File '{path}' written with root privileges ({len(content)} bytes)]"
            return f"[Sudo write error: {p.stderr.strip()}]"
        except Exception as e:
            return f"[Error writing with sudo: {e}]"
    except Exception as e:
        return f"[Error writing file: {e}]"

def tool_edit_file(path: Any, target_text: Any, replacement_text: Any) -> str:
    """Replaces specific text in any file, including Word .docx and system configs."""
    path = str(unwrap_tool_arg(path)).strip()
    target_text = str(unwrap_tool_arg(target_text))
    replacement_text = str(unwrap_tool_arg(replacement_text))
    expanded_path = os.path.expanduser(path)

    if not os.path.exists(expanded_path):
        candidates = [
            os.path.join(os.path.expanduser("~/Downloads"), path),
            os.path.join(os.path.expanduser("~"), path),
            os.path.join(os.path.expanduser("~/omarchy-virtual-paradise"), path),
            os.path.join(os.path.expanduser("~/.config/hypr"), path),
        ]
        for c in candidates:
            if os.path.exists(c):
                expanded_path = c
                path = c
                break

    if not os.path.exists(expanded_path):
        return f"[Error: File '{path}' not found]"

    if not prompt_preview_action("edit_file", f"Edit file '{expanded_path}'", f"'{target_text[:35]}' -> '{replacement_text[:35]}'"):
        return "[Action Cancelled: User declined edit operation]"

    # Support Word docx replacement
    if expanded_path.lower().endswith(".docx"):
        try:
            venv_py = os.path.expanduser("~/.local/lib/paradise-venv/bin/python3")
            py_bin = venv_py if os.path.exists(venv_py) else "python3"
            code = f"""
from docx import Document
doc = Document({repr(expanded_path)})
count = 0
for p in doc.paragraphs:
    if {repr(target_text)} in p.text:
        p.text = p.text.replace({repr(target_text)}, {repr(replacement_text)})
        count += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if {repr(target_text)} in p.text:
                    p.text = p.text.replace({repr(target_text)}, {repr(replacement_text)})
                    count += 1
doc.save({repr(expanded_path)})
print(f"Successfully replaced {{count}} instance(s) in Word document.")
"""
            out = subprocess.check_output([py_bin, "-c", code], text=True, timeout=10)
            return f"[Success editing Word document '{os.path.basename(path)}': {out.strip()}]"
        except Exception as e:
            return f"[Error editing Word document: {e}]"

    # Code / config / text file replacement
    try:
        with open(expanded_path, "r", encoding="utf-8") as f:
            content = f.read()
        if target_text not in content:
            return f"[Error: Text '{target_text[:60]}' not found in '{path}']"
        new_content = content.replace(target_text, replacement_text, 1)
        with open(expanded_path, "w", encoding="utf-8") as f:
            f.write(new_content)
        return f"[Success: File '{path}' updated ({len(target_text)} bytes replaced)]"
    except PermissionError:
        try:
            raw = subprocess.check_output(["sudo", "cat", expanded_path], text=True, timeout=5)
            if target_text not in raw:
                return f"[Error: Text '{target_text[:60]}' not found in '{path}']"
            new_content = raw.replace(target_text, replacement_text, 1)
            p = subprocess.run(["sudo", "tee", expanded_path], input=new_content, text=True, capture_output=True, timeout=10)
            if p.returncode == 0:
                return f"[Success: File '{path}' updated with root permissions]"
            return f"[Sudo edit error: {p.stderr.strip()}]"
        except Exception as e:
            return f"[Error editing with sudo: {e}]"
    except Exception as e:
        return f"[Error editing file: {e}]"

def tool_delete_file(path: Any) -> str:
    """Deletes any file or directory, auto-elevating to sudo if required."""
    path = str(unwrap_tool_arg(path)).strip()
    expanded_path = os.path.expanduser(path)

    if not os.path.exists(expanded_path):
        candidates = [
            os.path.join(os.path.expanduser("~/Downloads"), path),
            os.path.join(os.path.expanduser("~"), path),
        ]
        for c in candidates:
            if os.path.exists(c):
                expanded_path = c
                path = c
                break

    if not os.path.exists(expanded_path):
        return f"[Notice: Path '{path}' does not exist]"

    if not prompt_preview_action("delete_file", f"Delete path '{expanded_path}'"):
        return "[Action Cancelled: User declined deletion]"

    try:
        if os.path.isdir(expanded_path):
            shutil.rmtree(expanded_path)
        else:
            os.remove(expanded_path)
        return f"[Success: Deleted '{path}']"
    except PermissionError:
        return tool_execute_bash(f"sudo rm -rf '{expanded_path}'")
    except Exception as e:
        return f"[Error deleting: {e}]"

def tool_copy_file(source: Any, destination: Any) -> str:
    """Copies any file or directory across the filesystem."""
    src = str(unwrap_tool_arg(source)).strip()
    dst = str(unwrap_tool_arg(destination)).strip()
    expanded_src = os.path.expanduser(src)
    expanded_dst = os.path.expanduser(dst)

    if not os.path.exists(expanded_src):
        return f"[Error: Source '{src}' not found]"

    if not prompt_preview_action("copy_file", f"Copy '{expanded_src}' -> '{expanded_dst}'"):
        return "[Action Cancelled: User declined copy]"

    try:
        if os.path.isdir(expanded_src):
            shutil.copytree(expanded_src, expanded_dst, dirs_exist_ok=True)
        else:
            os.makedirs(os.path.dirname(os.path.abspath(expanded_dst)), exist_ok=True)
            shutil.copy2(expanded_src, expanded_dst)
        return f"[Success: Copied '{src}' to '{dst}']"
    except PermissionError:
        return tool_execute_bash(f"sudo cp -r '{expanded_src}' '{expanded_dst}'")
    except Exception as e:
        return f"[Error copying: {e}]"

def tool_move_file(source: Any, destination: Any) -> str:
    """Moves or renames any file or directory across the filesystem."""
    src = str(unwrap_tool_arg(source)).strip()
    dst = str(unwrap_tool_arg(destination)).strip()
    expanded_src = os.path.expanduser(src)
    expanded_dst = os.path.expanduser(dst)

    if not os.path.exists(expanded_src):
        return f"[Error: Source '{src}' not found]"

    if not prompt_preview_action("move_file", f"Move '{expanded_src}' -> '{expanded_dst}'"):
        return "[Action Cancelled: User declined move]"

    try:
        os.makedirs(os.path.dirname(os.path.abspath(expanded_dst)), exist_ok=True)
        shutil.move(expanded_src, expanded_dst)
        return f"[Success: Moved '{src}' to '{dst}']"
    except PermissionError:
        return tool_execute_bash(f"sudo mv '{expanded_src}' '{expanded_dst}'")
    except Exception as e:
        return f"[Error moving: {e}]"

def tool_get_system_health() -> str:
    """Collects comprehensive hardware and system health diagnostics."""
    report = []
    
    # 1. CPU & Load
    try:
        load = os.getloadavg()
        report.append(f"• CPU Load Average: 1m={load[0]:.2f}, 5m={load[1]:.2f}, 15m={load[2]:.2f}")
    except Exception:
        pass

    # 2. RAM & Swap
    mem = get_system_memory_info()
    report.append(f"• Memory: Total={mem['total']} GiB, Available={mem['available']} GiB, Free={mem['free']} GiB")

    # 3. Disk Space
    try:
        total, used, free = shutil.disk_usage("/")
        report.append(f"• Root Disk (/): Total={total // (2**30)} GiB, Used={used // (2**30)} GiB, Free={free // (2**30)} GiB")
    except Exception:
        pass

    # 4. Battery & Cooler Boost
    try:
        with open("/tmp/cooler_boost_state", "r") as f:
            boost_state = f.read().strip()
        report.append(f"• Hardware Cooler Boost: {boost_state.upper()}")
    except Exception:
        report.append("• Hardware Cooler Boost: AUTO")

    # 5. Top CPU Processes
    try:
        p = subprocess.run(["ps", "-eo", "pid,%cpu,%mem,comm", "--sort=-%cpu"], capture_output=True, text=True, timeout=2)
        top_procs = p.stdout.strip().split("\n")[1:5]
        report.append("• Top CPU Processes:\n  " + "\n  ".join(top_procs))
    except Exception:
        pass

    # 6. Failed System Services
    try:
        p = subprocess.run(["systemctl", "--failed", "--no-legend"], capture_output=True, text=True, timeout=2)
        failed = p.stdout.strip()
        if failed:
            report.append(f"• Failed Systemd Units:\n  {failed}")
        else:
            report.append("• Systemd Services: All services healthy (0 failed)")
    except Exception:
        pass

    return "\n".join(report)

def tool_load_skill(skill_name: Any) -> str:
    """Loads expert instructions for a specialized skill."""
    skill_name = str(unwrap_tool_arg(skill_name)).strip()
    skills = discover_skills()
    normalized = skill_name.lower().replace("_", "-").strip()
    matched = None
    for k, v in skills.items():
        if k.lower().replace("_", "-").strip() == normalized:
            matched = v
            break
    if not matched:
        for k, v in skills.items():
            if normalized in k.lower() or k.lower() in normalized:
                matched = v
                break
    if not matched:
        return f"[Error: Skill '{skill_name}' not found. Available skills: {', '.join(skills.keys())}]"

    try:
        with open(matched["path"], "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        lines = content.split("\n")
        filtered_lines = []
        in_frontmatter = False
        frontmatter_done = False
        for line in lines:
            if line.strip() == "---":
                if not in_frontmatter and not frontmatter_done:
                    in_frontmatter = True
                    continue
                elif in_frontmatter:
                    in_frontmatter = False
                    frontmatter_done = True
                    continue
            if not in_frontmatter:
                filtered_lines.append(line)
        body = "\n".join(filtered_lines).strip()
        return f"[Expert Skill Instructions: '{skill_name}']\n{body[:3500]}"
    except Exception as e:
        return f"[Error loading skill '{skill_name}': {e}]"

def tool_delete_skill(skill_name: Any) -> str:
    """Deletes and removes all directories of a specialized skill across all skill locations."""
    skill_name = str(unwrap_tool_arg(skill_name)).strip()
    if not skill_name:
        return "[Error: Missing skill name]"

    normalized = skill_name.lower().replace("_", "-").strip()
    deleted_paths = []

    for base in SKILL_DIRS:
        if not os.path.exists(base):
            continue
        try:
            items = os.listdir(base)
        except Exception:
            continue
        for item in items:
            item_norm = item.lower().replace("_", "-").strip()
            if item_norm == normalized or (normalized in item_norm and len(normalized) >= 4):
                full_p = os.path.join(base, item)
                if not prompt_preview_action("delete_skill", f"Delete skill folder '{full_p}'"):
                    return "[Action Cancelled: User declined skill deletion]"
                try:
                    if os.path.isdir(full_p):
                        shutil.rmtree(full_p)
                    else:
                        os.remove(full_p)
                    deleted_paths.append(full_p)
                except PermissionError:
                    subprocess.run(["sudo", "rm", "-rf", full_p], check=False)
                    deleted_paths.append(full_p)
                except Exception as e:
                    return f"[Error deleting skill '{skill_name}' at '{full_p}': {e}]"

    if deleted_paths:
        return f"[Success: Skill '{skill_name}' deleted from {len(deleted_paths)} location(s):\n  " + "\n  ".join(deleted_paths) + "]"
    return f"[Notice: Skill '{skill_name}' was not found in any active skill directories]"

def call_ollama_chat(messages: List[Dict[str, Any]], model: str, enable_tools: bool = True) -> Dict[str, Any]:
    """Sends chat completion request to local Ollama API with sliding context pruning and error resilience."""
    active_msgs = messages
    if len(messages) > 16:
        # Keep system prompt (index 0), initial user prompt (index 1), and recent conversation turns
        head = messages[:2]
        tail = messages[-12:]
        compacted_tail = []
        for m in tail:
            m_copy = dict(m)
            # Compact verbose tool outputs in historical turns to conserve token budget
            if m_copy.get("role") == "tool" and len(m_copy.get("content", "")) > 400:
                c = m_copy["content"]
                m_copy["content"] = c[:250] + "\n... [Output truncated for context optimization] ...\n" + c[-150:]
            compacted_tail.append(m_copy)
        active_msgs = head + compacted_tail

    payload: Dict[str, Any] = {
        "model": model,
        "messages": active_msgs,
        "stream": False,
        "options": {
            "temperature": 0.1,
            "num_thread": min(12, os.cpu_count() or 8),
            "num_ctx": 4096,
            "num_predict": 128 if enable_tools else 384
        }
    }
    if enable_tools:
        payload["tools"] = TOOLS_SPEC

    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_HOST}/api/chat",
        data=data,
        headers={"Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.URLError as e:
        return {"message": {"role": "assistant", "content": f"[Network Error communicating with Ollama: {e}]"}}
    except Exception as e:
        return {"message": {"role": "assistant", "content": f"[Error: {e}]"}}

def stream_ollama_synthesis(messages: List[Dict[str, Any]], model: str, lang: str = "vi") -> str:
    """Streams the assistant synthesis response directly to terminal in real-time for instant responsiveness."""
    active_msgs = messages
    if len(messages) > 16:
        head = messages[:2]
        tail = messages[-12:]
        compacted_tail = []
        for m in tail:
            m_copy = dict(m)
            if m_copy.get("role") == "tool" and len(m_copy.get("content", "")) > 400:
                c = m_copy["content"]
                m_copy["content"] = c[:250] + "\n... [Output truncated] ...\n" + c[-150:]
            compacted_tail.append(m_copy)
        active_msgs = head + compacted_tail

    payload = {
        "model": model,
        "messages": active_msgs,
        "stream": True,
        "options": {
            "temperature": 0.1,
            "num_thread": min(12, os.cpu_count() or 8),
            "num_ctx": 4096,
            "num_predict": 512
        }
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_HOST}/api/chat",
        data=data,
        headers={"Content-Type": "application/json"}
    )
    full_content = []
    console.print(f"\n[bold {PINK}]paradise-agent ❯[/] ", end="")
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            for line in resp:
                if not line.strip():
                    continue
                try:
                    chunk = json.loads(line.decode("utf-8"))
                    delta = chunk.get("message", {}).get("content", "")
                    if delta:
                        print(delta, end="", flush=True)
                        full_content.append(delta)
                except Exception:
                    pass
        print()
        return "".join(full_content)
    except Exception as e:
        console.print(f"\n[bold {RED}][Stream Error: {e}][/]")
        return "".join(full_content)

def check_ollama_alive() -> bool:
    """Checks if Ollama service is reachable."""
    try:
        req = urllib.request.Request(f"{OLLAMA_HOST}/api/tags")
        with urllib.request.urlopen(req, timeout=2) as resp:
            return resp.status == 200
    except Exception:
        return False

def get_banner_art() -> str:
    """Returns styled Cyberpunk ASCII banner with 45-degree diagonal gradient (Cyan -> Green -> Pink)."""
    raw_lines = [
        "██╗   ██╗██╗██████╗ ████████╗██╗   ██╗ █████╗ ██╗         ██████╗  █████╗ ██████╗  █████╗ ██████╗ ██╗███████╗███████╗",
        "██║   ██║██║██╔══██╗╚══██╔══╝██║   ██║██╔══██╗██║         ██╔══██╗██╔══██╗██╔══██╗██╔══██╗██╔══██╗██║██╔════╝██╔════╝",
        "██║   ██║██║██████╔╝   ██║   ██║   ██║███████║██║         ██████╔╝███████║██████╔╝███████║██║  ██║██║███████╗█████╗  ",
        "╚██╗ ██╔╝██║██╔══██╗   ██║   ██║   ██║██╔══██║██║         ██╔═══╝ ██╔══██║██╔══██╗██╔══██║██║  ██║██║╚════██║██╔══╝  ",
        " ╚████╔╝ ██║██║  ██║   ██║   ╚██████╔╝██║  ██║███████╗    ██║     ██║  ██║██║  ██║██║  ██║██████╔╝██║███████║███████╗",
        "  ╚═══╝  ╚═╝╚═╝  ╚═╝   ╚═╝    ╚═════╝ ╚═╝  ╚═╝╚══════╝    ╚═╝     ╚═╝  ╚═╝╚═╝  ╚═╝╚═╝  ╚═╝╚═════╝ ╚═╝╚══════╝╚══════╝"
    ]
    # Signature Virtual☆Paradise 40/20/40 Palette:
    # 40% Miku Cyan (#00f5d4) -> 20% Cyber Green (#00ff88) -> 40% Sakura Pink (#ffb7d5)
    c_cyan  = (0, 245, 212)
    c_green = (0, 255, 136)
    c_pink  = (255, 183, 213)

    h = len(raw_lines)
    w = max(len(l) for l in raw_lines)
    alpha = 2.0

    def lerp(c1, c2, p):
        p = max(0.0, min(1.0, p))
        return (
            int(c1[0] + (c2[0] - c1[0]) * p),
            int(c1[1] + (c2[1] - c1[1]) * p),
            int(c1[2] + (c2[2] - c1[2]) * p)
        )

    formatted_banner = []
    for r, line in enumerate(raw_lines):
        out = []
        for c, char in enumerate(line):
            if char == ' ':
                out.append(' ')
                continue
            t = (c + alpha * r) / (w + alpha * (h - 1))
            t = max(0.0, min(1.0, t))
            
            # True visual 40% Cyan / 20% Green / 40% Sakura Pink
            if t < 0.28:
                red, green, blue = c_cyan
            elif t < 0.38:
                red, green, blue = lerp(c_cyan, c_green, (t - 0.28) / 0.10)
            elif t < 0.62:
                red, green, blue = c_green
            elif t < 0.72:
                red, green, blue = lerp(c_green, c_pink, (t - 0.62) / 0.10)
            else:
                red, green, blue = c_pink

            out.append(f"\033[38;2;{red};{green};{blue}m{char}\033[0m")
        formatted_banner.append("  " + "".join(out))
    return "\n".join(formatted_banner)

def print_banner(model_name: str, reason: str = ""):
    """Renders Cyberpunk Authentic Diagonal Gradient Banner and TUI status bar."""
    art_text = get_banner_art()
    mem = get_system_memory_info()
    ram_str = f"{mem['available']} GiB Available / {mem['total']} GiB Total"
    mode_style = f"bold {GREEN}" if EXECUTION_MODE == "Auto-accept" else f"bold {YELLOW}"
    tier_desc = f" [dim]({reason})[/]" if reason else ""

    print(f"\n{art_text}\n")
    console.print(f"  [bold {MUTED}]Model:[/] [bold {GREEN}]{model_name}[/]{tier_desc}")
    console.print(f"  [bold {MUTED}]RAM:[/] [bold {CYAN}]{ram_str}[/]  |  [bold {MUTED}]Mode:[/] [{mode_style}]{EXECUTION_MODE}[/]")
    console.print(f"  [dim]Type [bold {YELLOW}]/[/] for commands, [bold {CYAN}]@[/] to tag files, [bold {YELLOW}]/help[/] for reference, [bold {YELLOW}]/exit[/] to quit.[/]")
    console.print(f"[dim {CYAN}]  {'─' * 95}[/]\n")

def print_thinking(thinking_text: str, title: str = "🧠 Diagnostic Reasoning"):
    """Displays agent's internal thought process in a dedicated TUI panel."""
    if not SHOW_THINKING or not thinking_text.strip():
        return
    console.print(Panel(
        thinking_text.strip(),
        title=f"[bold {PURPLE}]{title}[/]",
        border_style=PURPLE,
        box=box.ROUNDED,
        padding=(0, 1)
    ))

def print_help_table():
    """Displays commands in a clean TUI table."""
    table = Table(title="AGY Offline — Interactive Command Reference", border_style=CYAN, box=box.ROUNDED)
    table.add_column("Command", style=f"bold {YELLOW}", width=15)
    table.add_column("Description", style=f"{GREEN}")
    table.add_row("/resume", "Session manager: /resume [id|last|rename # name|delete #]")
    table.add_row("/plan", "Planning mode: /plan <task> (generate structured plan)")
    table.add_row("/goal", "Goal mode: /goal <objective> (autonomous execution)")
    table.add_row("/search", "Grep search across codebase: /search <query> [path]")
    table.add_row("/find", "Find files by glob pattern: /find <pattern> [dir]")
    table.add_row("/skills", "List all specialized expert skills available to the agent")
    table.add_row("/health", "Run immediate diagnostic health check (RAM, CPU, Disk, Services)")
    table.add_row("/mode", "Inspect or switch execution mode: /mode [auto|preview|toggle]")
    table.add_row("/model", "Inspect or switch local Ollama models: /model [1.5b|3b|7b]")
    table.add_row("/thinking", "Toggle real-time diagnostic reasoning visibility (ON / OFF)")
    table.add_row("/copy", "Copy last assistant response to clipboard")
    table.add_row("/diff", "View git diff or repository status: /diff [path]")
    table.add_row("/clear", "Clear active conversation memory context")
    table.add_row("/help", "Show this command reference table")
    table.add_row("/exit", "Exit session (Sayonara)")
    console.print(table)

def detect_language(text: str) -> str:
    """Detects whether user query is in Vietnamese or English."""
    vn_chars = "àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ"
    lower = text.lower()
    if any(c in vn_chars for c in lower):
        return "vi"
    vn_words = {"xin", "chao", "toi", "ban", "minh", "giup", "lay", "cua", "trong", "co", "gi", "may", "nay", "he", "thong", "la", "o", "va", "cho", "to", "trinh", "don", "mau", "van", "sao", "khong", "duoc", "xoa", "sua", "doc"}
    words = set(re.findall(r"\b\w+\b", lower))
    if words & vn_words:
        return "vi"
    return "en" if words else "vi"

def handle_user_turn(user_text: str, messages: List[Dict[str, Any]], model: str, session_id: Optional[str] = None):
    """Processes a multi-step user turn with reasoning extraction and tool execution."""
    global EXECUTION_MODE

    # Confirmation auto-dispatch: If user says "ok", "yes", etc., check if previous assistant turn proposed a command
    lower_u_strip = user_text.strip().lower().strip(".!?,")
    if lower_u_strip in ("ok", "yes", "y", "ừ", "uh", "đồng ý", "làm đi", "chạy đi", "xóa đi", "thực hiện đi", "sure", "tiến hành đi"):
        for prev in reversed(messages):
            if prev.get("role") == "assistant" and prev.get("content"):
                m_prev = re.search(r'```(?:bash|sh)?\s*\n(.*?)\n```', prev["content"], re.DOTALL)
                if m_prev:
                    proposed = m_prev.group(1).strip()
                    if proposed and not proposed.startswith("#"):
                        user_text = f"Execute command immediately: {proposed}"
                        break

    lang = detect_language(user_text)
    if lang == "vi":
        lang_directive = "[LANGUAGE MANDATE: The user explicitly wrote in VIETNAMESE. You MUST answer in natural, professional, polite VIETNAMESE.]"
    else:
        lang_directive = "[LANGUAGE MANDATE: DEFAULT LANGUAGE IS ENGLISH. Formulate your entire response in clear, fluent, professional ENGLISH.]"

    # Resolve and attach @file tags to prompt context
    at_tags = re.findall(r'@([^\s"\'`]+)', user_text)
    attached_files = []
    attached_descriptions = []

    for tag in at_tags:
        clean_tag = tag.strip(".,;:?!")
        resolved = None
        for cand in [
            clean_tag,
            os.path.expanduser(clean_tag),
            os.path.join(os.getcwd(), clean_tag),
            os.path.join(os.path.expanduser("~/Downloads"), clean_tag),
            os.path.join(os.path.expanduser("~"), clean_tag),
            os.path.join(os.path.expanduser("~/omarchy-virtual-paradise"), clean_tag),
            os.path.join(os.path.expanduser("~/.config/hypr"), clean_tag),
        ]:
            if os.path.exists(cand):
                resolved = os.path.abspath(cand)
                break
        if resolved and resolved not in attached_files:
            attached_files.append(resolved)
            file_body = tool_read_file(resolved, max_lines=150)
            attached_descriptions.append(f"\n[ATTACHED FILE: {resolved}]\n{file_body}\n[END ATTACHED FILE]")

    if attached_files:
        for af in attached_files:
            try:
                sz = os.path.getsize(af)
                sz_str = f"{sz/1024:.1f} KB" if sz < 1024*1024 else f"{sz/(1024*1024):.1f} MB"
            except Exception:
                sz_str = ""
            console.print(Panel(
                f"[bold {CYAN}]📎 Attached File Context:[/] [bold {PINK}]{af}[/] [dim]({sz_str})[/]",
                border_style=CYAN,
                box=box.ROUNDED,
                padding=(0, 1)
            ))
        user_text += "\n" + "\n".join(attached_descriptions)

    augmented_user_text = f"{user_text}\n\n{lang_directive}"
    messages.append({"role": "user", "content": augmented_user_text})
    if session_id:
        save_session(session_id, messages, model)

    max_steps = 8
    step = 0

    while step < max_steps:
        step += 1
        console.print(f"[{MUTED}]⚡ Analyzing intent locally...[/]", end="\r")

        start_time = time.time()
        try:
            res = call_ollama_chat(messages, model, enable_tools=(step == 1))
        except Exception as e:
            console.print(f"\n[bold {RED}][Model Error]:[/] {e}")
            break

        elapsed_ms = int((time.time() - start_time) * 1000)
        console.print(" " * 45, end="\r")

        msg = res.get("message", {})
        content = msg.get("content", "")
        tool_calls = msg.get("tool_calls", [])

        # Extract <think>...</think> blocks if model generated internal chain-of-thought
        extracted_think = ""
        if "<think>" in content and "</think>" in content:
            parts = content.split("<think>", 1)[1].split("</think>", 1)
            extracted_think = parts[0].strip()
            content = (content.split("<think>", 1)[0] + parts[1]).strip()

        # Parse inline tool calls if model emitted them in content
        if not tool_calls and content:
            trimmed = content.strip()

            # Markdown bash block extractor
            m_bash = re.search(r'```(?:bash|sh)?\s*\n(.*?)\n```', content, re.DOTALL)
            if m_bash:
                cmd_str = m_bash.group(1).strip()
                if cmd_str and not cmd_str.startswith("#"):
                    tool_calls = [{"function": {"name": "execute_bash", "arguments": {"command": cmd_str}}}]
                    content = ""

            if not tool_calls and "```json" in trimmed:
                try:
                    trimmed = trimmed.split("```json")[1].split("```")[0].strip()
                except Exception:
                    pass
            elif not tool_calls and "```" in trimmed:
                try:
                    trimmed = trimmed.split("```")[1].split("```")[0].strip()
                except Exception:
                    pass

            for line in trimmed.split("\n"):
                line = line.strip()
                if line.startswith("{") and line.endswith("}") and "name" in line:
                    try:
                        parsed = json.loads(line)
                        raw_name = str(parsed.get("name", "")).strip()
                        raw_args = parsed.get("arguments") or parsed.get("parameters", {})
                        c_name, c_args = normalize_tool_call(raw_name, raw_args if isinstance(raw_args, dict) else {})
                        if c_name in ALL_VALID_TOOLS:
                            tool_calls.append({"function": {"name": c_name, "arguments": c_args}})
                    except Exception:
                        pass

            if not tool_calls and trimmed.startswith("{") and "name" in trimmed:
                try:
                    parsed = json.loads(trimmed)
                    raw_name = str(parsed.get("name", "")).strip()
                    raw_args = parsed.get("arguments") or parsed.get("parameters", {})
                    c_name, c_args = normalize_tool_call(raw_name, raw_args if isinstance(raw_args, dict) else {})
                    if c_name in ALL_VALID_TOOLS:
                        tool_calls = [{"function": {"name": c_name, "arguments": c_args}}]
                    elif raw_name and not raw_name.startswith("{"):
                        tool_calls = [{"function": {"name": "execute_bash", "arguments": {"command": raw_name}}}]
                except Exception:
                    pass

            if tool_calls:
                content = ""

        # Safeguard: Never execute more than 1 tool call per step to prevent hallucinations
        if len(tool_calls) > 1:
            tool_calls = tool_calls[:1]

        if not msg.get("tool_calls") and tool_calls:
            msg["tool_calls"] = tool_calls
            msg["content"] = ""
        elif msg.get("tool_calls") and len(msg["tool_calls"]) > 1:
            msg["tool_calls"] = msg["tool_calls"][:1]
            tool_calls = msg["tool_calls"]

        # Proactive Grounding & Anti-Hallucination Guard on step 1
        lower_u = user_text.lower()
        user_home = os.path.expanduser("~")

        # 0. Delete skill guard (overrides hallucinated load_skill or bash delete_skill)
        if step == 1 and any(w in lower_u for w in ["xóa", "delete", "remove", "gỡ", "hủy"]) and (any(w in lower_u for w in ["skill", "kỹ năng"]) or any(sk.lower() in lower_u for sk in discover_skills().keys())):
            m_sk = re.search(r'(?:xóa|gỡ|hủy|remove|delete)\s+(?:skill|kỹ năng)?\s*([\w-]+)', lower_u)
            if not m_sk:
                m_sk = re.search(r'skill\s+([\w-]+)\s+(?:xóa|gỡ|hủy|remove|delete)', lower_u)
            target_sk = m_sk.group(1).strip() if m_sk else None
            if not target_sk or target_sk in ("đi", "bỏ", "file", "tệp"):
                for sk in discover_skills().keys():
                    if sk.lower() in lower_u:
                        target_sk = sk
                        break
            if not target_sk and "windows" in lower_u:
                target_sk = "windows"
            if target_sk:
                tool_calls = [{
                    "function": {
                        "name": "delete_skill",
                        "arguments": {"skill_name": target_sk}
                    }
                }]
                msg["tool_calls"] = tool_calls
                msg["content"] = ""
                content = ""

        elif not tool_calls and step == 1:

            # 1. Format document
            if any(w in lower_u for w in ["format", "chuẩn hóa", "định dạng"]) and any(w in lower_u for w in ["tờ trình", "to_trinh", "docx", "word", "tài liệu"]):
                tool_calls = [{
                    "function": {
                        "name": "execute_bash",
                        "arguments": {"command": f"{user_home}/.local/lib/paradise-venv/bin/python3 {user_home}/.local/bin/format-docx-vn.py {user_home}/Downloads"}
                    }
                }]
            # 2. Template tờ trình
            elif any(w in lower_u for w in ["mẫu", "template", "xin mẫu", "lấy file mẫu", "tải mẫu"]) and any(w in lower_u for w in ["tờ trình", "to_trinh", "công văn"]):
                tool_calls = [{
                    "function": {
                        "name": "execute_bash",
                        "arguments": {"command": f"cp {user_home}/Windows/skills/vn-officecli/templates/to_trinh_mau.docx {user_home}/Downloads/ && ls -la {user_home}/Downloads/"}
                    }
                }]
            # 3. Read / view file content
            elif any(w in lower_u for w in ["nội dung", "xem file", "đọc file", "mở file", "trong file", "file này", "read file", "view file", "cat "]):
                target_file = None
                m_f = re.search(r'[\w.-]+\.(docx|conf|toml|ini|json|txt|md|py|sh|lua|log|yaml|yml)', lower_u)
                if m_f:
                    target_file = m_f.group(0)
                elif any(w in lower_u for w in ["tờ trình", "to_trinh"]):
                    target_file = f"{user_home}/Downloads/to_trinh_mau.docx"
                elif any(w in lower_u for w in ["này", "this"]):
                    for prev in reversed(messages):
                        c = prev.get("content", "")
                        m_prev = re.search(r'[\w.-]+\.(docx|conf|toml|ini|json|txt|md|py|sh|lua|log|yaml|yml)', c)
                        if m_prev:
                            target_file = m_prev.group(0)
                            break
                if not target_file:
                    target_file = f"{user_home}/Downloads/to_trinh_mau.docx"

                tool_calls = [{
                    "function": {
                        "name": "read_file",
                        "arguments": {"path": target_file}
                    }
                }]
            # 4. Downloads directory
            elif any(w in lower_u for w in ["download", "downloads", "tải về"]) and any(w in lower_u for w in ["file", "tệp", "gì", "mục", "thư mục", "xem", "danh sách", "có", "what", "list", "show"]):
                tool_calls = [{
                    "function": {
                        "name": "execute_bash",
                        "arguments": {"command": f"ls -la {user_home}/Downloads"}
                    }
                }]
            # 5. Windows directory
            elif any(w in lower_u for w in ["window", "windows"]) and any(w in lower_u for w in ["file", "tệp", "gì", "mục", "thư mục", "xem", "danh sách", "có", "what", "list", "show"]):
                tool_calls = [{
                    "function": {
                        "name": "execute_bash",
                        "arguments": {"command": f"ls -la {user_home}/Windows"}
                    }
                }]
            # 6. Omarchy theme / dotfiles repo directory
            elif any(w in lower_u for w in ["omarchy-virtual-paradise", "virtual-paradise", "theme folder"]) and any(w in lower_u for w in ["file", "tệp", "gì", "mục", "thư mục", "xem", "danh sách", "có", "what", "list", "show"]):
                tool_calls = [{
                    "function": {
                        "name": "execute_bash",
                        "arguments": {"command": f"ls -la {user_home}/omarchy-virtual-paradise"}
                    }
                }]
            # 7. System health / performance
            elif any(w in lower_u for w in ["chậm", "lag", "đơ", "sức khỏe", "tình trạng máy", "kiểm tra máy", "pin", "ram", "cpu", "slow", "battery", "health"]):
                tool_calls = [{
                    "function": {
                        "name": "get_system_health",
                        "arguments": {}
                    }
                }]
            # 8. Delete / remove skills or files
            elif any(w in lower_u for w in ["xóa", "delete", "remove", "gỡ", "hủy"]):
                if any(w in lower_u for w in ["skill", "kỹ năng"]) or any(sk.lower() in lower_u for sk in discover_skills().keys()):
                    m_sk = re.search(r'(?:xóa|gỡ|hủy|remove|delete)\s+(?:skill|kỹ năng)?\s*([\w-]+)', lower_u)
                    if not m_sk:
                        m_sk = re.search(r'skill\s+([\w-]+)\s+(?:xóa|gỡ|hủy|remove|delete)', lower_u)
                    target_sk = m_sk.group(1).strip() if m_sk else None
                    if not target_sk or target_sk in ("đi", "bỏ", "file", "tệp"):
                        for sk in discover_skills().keys():
                            if sk.lower() in lower_u:
                                target_sk = sk
                                break
                    if not target_sk and "windows" in lower_u:
                        target_sk = "windows"
                    if target_sk:
                        tool_calls = [{
                            "function": {
                                "name": "delete_skill",
                                "arguments": {"skill_name": target_sk}
                            }
                        }]
                elif any(w in lower_u for w in ["formatted", "chuẩn hóa"]):
                    tool_calls = [{
                        "function": {
                            "name": "delete_file",
                            "arguments": {"path": f"{user_home}/Downloads/to_trinh_mau_formatted.docx"}
                        }
                    }]
                elif any(w in lower_u for w in ["to_trinh", "tờ trình", "mẫu"]):
                    tool_calls = [{
                        "function": {
                            "name": "delete_file",
                            "arguments": {"path": f"{user_home}/Downloads/to_trinh_mau.docx"}
                        }
                    }]
                else:
                    m_file = re.search(r'[\w.-]+\.\w+', lower_u)
                    if m_file:
                        fn = m_file.group(0)
                        tool_calls = [{
                            "function": {
                                "name": "delete_file",
                                "arguments": {"path": f"{user_home}/Downloads/{fn}"}
                            }
                        }]
            # 9. Cooler Boost / Fan control
            elif any(w in lower_u for w in ["cooler boost", "coolerboost", "quạt gió", "quạt tản nhiệt", "fan boost", "bật quạt", "tắt quạt"]):
                if any(w in lower_u for w in ["tắt", "stop", "off", "disable"]):
                    tool_calls = [{
                        "function": {
                            "name": "execute_bash",
                            "arguments": {"command": f"{user_home}/.local/bin/toggle_cooler_boost.sh off"}
                        }
                    }]
                else:
                    tool_calls = [{
                        "function": {
                            "name": "execute_bash",
                            "arguments": {"command": f"{user_home}/.local/bin/toggle_cooler_boost.sh on"}
                        }
                    }]

            if tool_calls:
                msg["tool_calls"] = tool_calls
                msg["content"] = ""
                content = ""

        # Print reasoning / thinking process if available or display telemetry
        if SHOW_THINKING:
            diagnostic_reasoning = ""
            if extracted_think:
                diagnostic_reasoning = f"[bold {PURPLE}]Internal Reasoning Stream:[/]\n{extracted_think}\n\n"
            if tool_calls:
                fn_intent = tool_calls[0].get("function", {}).get("name", "action")
                diagnostic_reasoning += (
                    f"• [bold {CYAN}]Step {step} Analysis:[/] Intent verified | Target tool: [bold {YELLOW}]{fn_intent}[/]\n"
                    f"• [bold {CYAN}]Model Decision Time:[/] {elapsed_ms} ms | Lang: [bold]{lang.upper()}[/] | Mode: [bold]{EXECUTION_MODE}[/]"
                )
            elif content:
                diagnostic_reasoning += (
                    f"• [bold {CYAN}]Synthesis Step:[/] Formulation completed in {elapsed_ms} ms | Language: [bold]{lang.upper()}[/]"
                )
            if diagnostic_reasoning.strip():
                print_thinking(diagnostic_reasoning)

        if content:
            console.print(f"\n[bold {PINK}]paradise-agent ❯[/]")
            console.print(Markdown(content))

        messages.append(msg)

        if not tool_calls:
            # ReAct turn completed
            break

        # Execute requested tools
        for tc in tool_calls:
            fn = tc.get("function", {})
            raw_fn_name = str(fn.get("name", "")).strip()
            raw_fn_args = fn.get("arguments", {})
            if isinstance(raw_fn_args, str):
                try:
                    raw_fn_args = json.loads(raw_fn_args)
                except Exception:
                    raw_fn_args = {"command": raw_fn_args}
            if not isinstance(raw_fn_args, dict):
                raw_fn_args = {"command": str(raw_fn_args)}

            fn_name, fn_args = normalize_tool_call(raw_fn_name, raw_fn_args)

            # Display tool invocation TUI
            console.print(Panel(
                f"[bold {CYAN}]{fn_name}[/]([dim]{json.dumps(fn_args, ensure_ascii=False)}[/])",
                title=f"[bold {YELLOW}]⚙ Tool Execution[/]",
                border_style=YELLOW,
                box=box.ROUNDED,
                padding=(0, 1)
            ))

            tool_start = time.time()
            tool_output = ""
            try:
                if fn_name == "execute_bash":
                    cmd = fn_args.get("command", "")
                    tool_output = tool_execute_bash(cmd)
                elif fn_name == "read_file":
                    p = fn_args.get("path", "")
                    sl = fn_args.get("start_line")
                    el = fn_args.get("end_line")
                    ml = fn_args.get("max_lines", 200)
                    tool_output = tool_read_file(p, start_line=sl, end_line=el, max_lines=ml)
                elif fn_name == "grep_search":
                    q = fn_args.get("query", "")
                    sp = fn_args.get("search_path", ".")
                    ci = fn_args.get("case_insensitive", True)
                    tool_output = tool_grep_search(q, sp, case_insensitive=ci)
                elif fn_name == "find_by_name":
                    pat = fn_args.get("pattern", "")
                    sd = fn_args.get("search_directory", ".")
                    tool_output = tool_find_by_name(pat, sd)
                elif fn_name == "list_dir":
                    dp = fn_args.get("directory_path", ".")
                    tool_output = tool_list_dir(dp)
                elif fn_name == "write_file":
                    p = fn_args.get("path", "")
                    c = fn_args.get("content", "")
                    tool_output = tool_write_file(p, c)
                elif fn_name == "edit_file":
                    p = fn_args.get("path", "")
                    t = fn_args.get("target_text", "")
                    r = fn_args.get("replacement_text", "")
                    tool_output = tool_edit_file(p, t, r)
                elif fn_name == "delete_file":
                    p = fn_args.get("path", "")
                    tool_output = tool_delete_file(p)
                elif fn_name == "copy_file":
                    s = fn_args.get("source", "")
                    d = fn_args.get("destination", "")
                    tool_output = tool_copy_file(s, d)
                elif fn_name == "move_file":
                    s = fn_args.get("source", "")
                    d = fn_args.get("destination", "")
                    tool_output = tool_move_file(s, d)
                elif fn_name == "get_system_health":
                    tool_output = tool_get_system_health()
                elif fn_name == "load_skill":
                    s_name = fn_args.get("skill_name", "")
                    tool_output = tool_load_skill(s_name)
                elif fn_name == "delete_skill":
                    s_name = fn_args.get("skill_name", "")
                    tool_output = tool_delete_skill(s_name)
                else:
                    tool_output = f"[Unknown tool: {fn_name}]"
            except Exception as exc:
                tool_output = f"[Tool Error in {fn_name}: {exc}]"

            tool_time_ms = int((time.time() - tool_start) * 1000)

            # Display concise preview of tool output
            preview = tool_output.strip().split("\n")
            preview_str = "\n".join(preview[:8])
            if len(preview) > 8:
                preview_str += f"\n... ({len(preview) - 8} more lines)"

            console.print(Panel(
                preview_str,
                title=f"[bold {GREEN}]📄 Tool Result ({tool_time_ms} ms)[/]",
                border_style=GREEN,
                box=box.ROUNDED,
                padding=(0, 1)
            ))

            tool_content_for_model = tool_output
            if lang == "vi":
                tool_content_for_model += "\n[SYSTEM MANDATE: You MUST reply to the user completely in natural VIETNAMESE. Do NOT reply in French, English, or any other language.]"
            else:
                tool_content_for_model += "\n[SYSTEM MANDATE: You MUST reply to the user completely in ENGLISH.]"

            messages.append({
                "role": "tool",
                "name": fn_name,
                "content": tool_content_for_model
            })

        # Directly stream the assistant synthesis response to the terminal in real-time!
        synth_start = time.time()
        synth_content = stream_ollama_synthesis(messages, model, lang=lang)
        synth_ms = int((time.time() - synth_start) * 1000)
        if SHOW_THINKING and synth_content:
            print_thinking(f"• [bold {CYAN}]Synthesis Stream:[/] Delivered in {synth_ms} ms | Language: [bold]{lang.upper()}[/]")
        messages.append({
            "role": "assistant",
            "content": synth_content
        })
        break

    if session_id:
        save_session(session_id, messages, model)

def agent_loop(initial_prompt: Optional[str] = None, requested_model: Optional[str] = None, resume_session: Optional[str] = None):
    """Main interactive REPL loop with session persistence and /resume support."""
    global EXECUTION_MODE, SHOW_THINKING

    model, reason = detect_optimal_model(requested_model)
    session_id = f"session_{time.strftime('%Y%m%d_%H%M%S')}"

    if not check_ollama_alive():
        console.print(f"[bold {YELLOW}][!] Ollama service is not responding. Attempting to start...[/]")
        try:
            subprocess.run(["systemctl", "start", "ollama.service"], check=False)
            time.sleep(1.5)
        except Exception:
            pass
        if not check_ollama_alive():
            console.print(f"[bold {RED}][Error] Could not connect to Ollama. Please run 'systemctl start ollama'.[/]")
            sys.exit(1)

    messages = [
        {"role": "system", "content": build_system_prompt()}
    ]

    if resume_session:
        res = load_session(resume_session)
        if res:
            session_id, messages, s_model = res
            if not requested_model:
                model = s_model
            print_banner(model, reason)
            render_restored_session(messages)
        else:
            print_banner(model, reason)
            console.print(f"[bold {RED}]Session '{resume_session}' not found. Starting fresh session.[/]")
    else:
        print_banner(model, reason)

    if initial_prompt:
        handle_user_turn(initial_prompt, messages, model, session_id=session_id)
        return

    prompt_session = None
    if HAS_PROMPT_TOOLKIT:
        prompt_style = Style.from_dict({
            'completion-menu.completion': 'bg:#12151e #00f5d4',
            'completion-menu.completion.current': 'bg:#00f5d4 #0f111a bold',
            'completion-menu.meta.completion': 'bg:#1a1c26 #ffb7d5',
            'completion-menu.meta.completion.current': 'bg:#00f5d4 #1a1c26 bold',
            'scrollbar.background': 'bg:#0f111a',
            'scrollbar.button': 'bg:#00ff88',
        })
        history_file = os.path.expanduser("~/.local/share/paradise-agent/history")
        os.makedirs(os.path.dirname(history_file), exist_ok=True)
        prompt_session = PromptSession(
            history=FileHistory(history_file),
            completer=SlashAndFileCompleter(),
            complete_while_typing=True,
            style=prompt_style
        )

    while True:
        try:
            if prompt_session:
                user_input = prompt_session.prompt(ANSI("\033[38;2;0;245;212m\033[1myou ❯\033[0m ")).strip()
            else:
                user_input = input(f"\n\033[38;2;0;245;212m\033[1myou ❯\033[0m ").strip()
        except (KeyboardInterrupt, EOFError):
            console.print(f"\n[bold {PINK}]Farewell from Virtual☆Paradise! // Sayonara.[/]")
            break

        if not user_input:
            continue

        if user_input.startswith("/"):
            cmd = user_input.lower().strip()
            if cmd in ("/exit", "/quit", "/q"):
                console.print(f"[bold {PINK}]Farewell from Virtual☆Paradise! // Sayonara.[/]")
                break
            elif cmd == "/help":
                print_help_table()
                continue
            elif cmd in ("/skills", "/skill"):
                skills_map = discover_skills()
                table = Table(title=f"Discovered Skills ({len(skills_map)} available)", border_style=CYAN, box=box.ROUNDED)
                table.add_column("Skill Name", style=f"bold {CYAN}")
                table.add_column("Description", style=f"{GREEN}")
                for k, v in skills_map.items():
                    table.add_row(k, v["description"][:100] + "...")
                console.print(table)
                continue
            elif cmd == "/health":
                console.print(Panel(tool_get_system_health(), title=f"[bold {CYAN}]Hardware & System Health[/]", border_style=CYAN, box=box.ROUNDED))
                continue
            elif cmd == "/clear":
                session_id = f"session_{time.strftime('%Y%m%d_%H%M%S')}"
                messages = [{"role": "system", "content": build_system_prompt()}]
                console.print(f"[bold {GREEN}]Active conversation memory cleared (New session started).[/]")
                continue
            elif cmd.startswith("/resume"):
                parts = user_input.split(maxsplit=1)
                arg_val = parts[1].strip() if len(parts) > 1 else ""
                res = handle_resume_command(arg_val)
                if res:
                    session_id, messages, s_model = res
                    model = s_model
                    render_restored_session(messages)
                continue
            elif cmd.startswith("/plan"):
                parts = user_input.split(maxsplit=1)
                task_content = parts[1].strip() if len(parts) > 1 else ""
                if not task_content:
                    try:
                        task_content = input(f"\033[38;2;255;224;102mEnter task/architecture to plan: \033[0m").strip()
                    except (KeyboardInterrupt, EOFError):
                        continue
                if not task_content:
                    console.print(f"[dim]Planning cancelled.[/]")
                    continue
                plan_prompt = (
                    f"[PLANNING DIRECTIVE]: You are in deep architectural planning mode. "
                    f"Formulate a structured, step-by-step diagnostic and execution plan with clear numbered phases "
                    f"and verification criteria for: {task_content}"
                )
                handle_user_turn(plan_prompt, messages, model, session_id=session_id)
                continue
            elif cmd.startswith("/goal"):
                parts = user_input.split(maxsplit=1)
                goal_content = parts[1].strip() if len(parts) > 1 else ""
                if not goal_content:
                    try:
                        goal_content = input(f"\033[38;2;255;224;102mEnter goal/objective to achieve: \033[0m").strip()
                    except (KeyboardInterrupt, EOFError):
                        continue
                if not goal_content:
                    console.print(f"[dim]Goal cancelled.[/]")
                    continue
                goal_prompt = (
                    f"[AUTONOMOUS GOAL DIRECTIVE]: Execute all necessary diagnostic, inspection, and coding steps "
                    f"autonomously without stopping until this goal is fully accomplished and verified: {goal_content}"
                )
                handle_user_turn(goal_prompt, messages, model, session_id=session_id)
                continue
            elif cmd in ("/copy", "/cp"):
                last_resp = ""
                for m in reversed(messages):
                    if m.get("role") == "assistant" and m.get("content"):
                        last_resp = m["content"]
                        break
                if last_resp:
                    try:
                        p = subprocess.Popen(["wl-copy"], stdin=subprocess.PIPE, text=True)
                        p.communicate(input=last_resp, timeout=3)
                        console.print(f"[bold {GREEN}]✔ Copied last assistant response to Wayland clipboard.[/]")
                    except Exception as e:
                        console.print(f"[bold {RED}]Failed to copy to clipboard: {e}[/]")
                else:
                    console.print(f"[bold {YELLOW}]No assistant response to copy yet.[/]")
                continue
            elif cmd.startswith("/diff"):
                parts = user_input.split(maxsplit=1)
                target_repo = parts[1].strip() if len(parts) > 1 else "."
                try:
                    p = subprocess.run(["git", "diff", "--stat", "-p"], cwd=os.path.expanduser(target_repo), capture_output=True, text=True, timeout=5)
                    out = p.stdout.strip()
                    if not out:
                        p_st = subprocess.run(["git", "status", "-s"], cwd=os.path.expanduser(target_repo), capture_output=True, text=True, timeout=5)
                        out = p_st.stdout.strip() or "No local changes in repository."
                    console.print(Panel(out[:4000], title=f"[bold {CYAN}]Git Diff / Status ({target_repo})[/]", border_style=CYAN, box=box.ROUNDED))
                except Exception as e:
                    console.print(f"[bold {RED}]Git diff error: {e}[/]")
                continue
            elif cmd.startswith("/thinking"):
                SHOW_THINKING = not SHOW_THINKING
                state = f"[bold {GREEN}]ON[/]" if SHOW_THINKING else f"[bold {RED}]OFF[/]"
                console.print(f"Diagnostic thinking display is now: {state}")
                continue
            elif cmd == "/model" or cmd.startswith("/model "):
                parts = user_input.split()
                if len(parts) > 1:
                    req_m = parts[1].strip()
                    installed = get_installed_models()
                    matched_model = req_m
                    for inst in installed:
                        if req_m.lower() in inst.lower():
                            matched_model = inst
                            break
                    model = matched_model
                    console.print(f"[bold {GREEN}]Active model switched to:[/] [bold {CYAN}]{model}[/]")
                else:
                    installed = get_installed_models()
                    mem = get_system_memory_info()
                    table = Table(title="System Memory & Installed Models", border_style=CYAN, box=box.ROUNDED)
                    table.add_column("Model Name", style=f"bold {CYAN}")
                    table.add_column("Status", style=f"{GREEN}")
                    for m in installed:
                        status = "✔ Active" if m == model else "Available"
                        table.add_row(m, status)
                    console.print(f"RAM Available: [bold {CYAN}]{mem['available']} GiB[/] / {mem['total']} GiB Total")
                    console.print(table)
                    console.print(f"[dim]Tip: Use '/model <name>' (e.g. '/model 7b' or '/model 3b') to switch models on the fly.[/]")
                continue
            elif cmd == "/mode" or cmd.startswith("/mode "):
                parts = user_input.split()
                if len(parts) > 1:
                    target = parts[1].lower().strip()
                    if target in ("preview", "p"):
                        EXECUTION_MODE = "Preview"
                    elif target in ("auto", "auto-accept", "a", "yolo"):
                        EXECUTION_MODE = "Auto-accept"
                    elif target == "toggle":
                        EXECUTION_MODE = "Preview" if EXECUTION_MODE == "Auto-accept" else "Auto-accept"
                    console.print(f"[bold {GREEN}]Execution Mode set to:[/] [bold {YELLOW}]{EXECUTION_MODE}[/]")
                else:
                    mode_table = Table(title="Execution Mode Configuration", border_style=YELLOW, box=box.ROUNDED)
                    mode_table.add_column("Property", style="bold")
                    mode_table.add_column("Setting", style=f"bold {GREEN}")
                    mode_table.add_row("Current Mode", EXECUTION_MODE)
                    mode_table.add_row("Auto-accept", "Autonomous tool execution without prompts (Default)")
                    mode_table.add_row("Preview", "Interactive approval prompt before any tool action")
                    mode_table.add_row("Usage", "/mode auto  |  /mode preview  |  /mode toggle")
                    console.print(mode_table)
                continue
            elif cmd.startswith("/search"):
                parts = user_input.split(maxsplit=2)
                if len(parts) > 1:
                    query = parts[1]
                    path = parts[2] if len(parts) > 2 else "."
                    console.print(Panel(tool_grep_search(query, path), title=f"[bold {CYAN}]Grep Search: '{query}' in '{path}'[/]", border_style=CYAN, box=box.ROUNDED))
                else:
                    console.print(f"[bold {YELLOW}]Usage: /search <pattern> [path][/]")
                continue
            elif cmd.startswith("/find"):
                parts = user_input.split(maxsplit=2)
                if len(parts) > 1:
                    pattern = parts[1]
                    dir_path = parts[2] if len(parts) > 2 else "."
                    console.print(Panel(tool_find_by_name(pattern, dir_path), title=f"[bold {CYAN}]Find Files: '{pattern}' in '{dir_path}'[/]", border_style=CYAN, box=box.ROUNDED))
                else:
                    console.print(f"[bold {YELLOW}]Usage: /find <glob_pattern> [dir][/]")
                continue

        handle_user_turn(user_input, messages, model, session_id=session_id)

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Virtual☆Paradise Autonomous Diagnostic & Coding Agent")
    parser.add_argument("prompt", nargs="?", help="Optional initial instruction to execute")
    parser.add_argument("--preview", "-p", action="store_true", help="Start in Preview mode (prompts before executing actions)")
    parser.add_argument("--yolo", "-y", "--auto-accept", action="store_true", help="Start in Auto-accept mode (default)")
    parser.add_argument("--model", "-m", help="Override Ollama model name")
    parser.add_argument("--resume", "-r", nargs="?", const="last", help="Resume past conversation session (specify session ID, index number, or 'last')")
    parser.add_argument("--plan", help="Run planning mode with the specified task")
    parser.add_argument("--goal", help="Run autonomous goal execution mode with the specified objective")
    parser.add_argument("--diagnose", "-d", action="store_true", help="Run system health diagnosis immediately")

    args = parser.parse_args()

    global EXECUTION_MODE
    if args.preview:
        EXECUTION_MODE = "Preview"
    else:
        EXECUTION_MODE = "Auto-accept"

    if args.diagnose:
        console.print(Panel(tool_get_system_health(), title=f"[bold {CYAN}]System Health Diagnostics[/]", border_style=CYAN, box=box.ROUNDED))
        sys.exit(0)

    initial = args.prompt
    if args.plan:
        initial = f"[PLANNING DIRECTIVE]: Formulate a structured step-by-step diagnostic and execution plan with clear numbered phases and verification criteria for: {args.plan}"
    elif args.goal:
        initial = f"[AUTONOMOUS GOAL DIRECTIVE]: Execute all necessary diagnostic, inspection, and coding steps autonomously without stopping until this goal is fully accomplished and verified: {args.goal}"

    agent_loop(initial_prompt=initial, requested_model=args.model, resume_session=args.resume)

if __name__ == "__main__":
    main()
